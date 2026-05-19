#!/usr/bin/env python3
"""
Merge supervisor revisions into ScribeFlow_FINAL.docx while keeping all figures (WMF).
Output: ScribeFlow_REVISED.docx
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "ScribeFlow_FINAL.docx"
OUT = ROOT / "ScribeFlow_REVISED.docx"

PROD_URL = "https://scribeflow-two.vercel.app"


def replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    text = paragraph.text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    return True


def replace_all(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, old, new):
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_paragraph(p, old, new):
                        n += 1
    return n


def insert_after_heading(doc: Document, heading_prefix: str, lines: list[str]) -> None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(heading_prefix):
            ref = p._element
            parent = ref.getparent()
            idx = parent.index(ref)
            for offset, line in enumerate(lines, start=1):
                new_p = doc.add_paragraph(line)
                parent.insert(idx + offset, new_p._element)
            return


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Missing {SRC}")

    shutil.copy(SRC, OUT)
    doc = Document(OUT)

    replacements = [
        (
            "production-ready system",
            "deployable cloud-native prototype",
        ),
        (
            "role-aware workspace collaboration features",
            "workspace-scoped ACL (owner and collaborator)",
        ),
        (
            "This predicate is consistent with the RBAC model described by Sandhu et al. [8], with workspace ownership and collaborator membership serving as the two principal roles.",
            "This predicate implements workspace-scoped access control (ACL): the owner (workspace_owner_id) and collaborators (collaborators table). It is not hierarchical RBAC as defined by Sandhu et al. [8].",
        ),
        (
            "ScribeFlow adopts a simplified two-role model (owner and collaborator) consistent with the core RBAC specification, extended with workspace-scoped evaluation to prevent cross-tenant access.",
            "ScribeFlow adopts a two-principal ACL (owner and collaborator) with workspace-scoped evaluation to prevent cross-tenant access. Hierarchical RBAC is identified as future work.",
        ),
        (
            "Role-based Access Control",
            "Workspace ACL (owner / collaborator)",
        ),
        (
            "Each metric was averaged over 50 independent requests under a simulated concurrent load of 10 users, implemented using the k6 load-testing framework.",
            "Load tests used constant-arrival-rate scripts (benchmarks/k6/load.js and benchmarks/run-load-test.mjs) at 80–500 requests per second (RPS), reporting p50, p95, p99, mean, and standard deviation (Tables IV and VI).",
        ),
        (
            "Network: 100 Mbps symmetric fiber (measured RTT to Vercel edge: ~12 ms)",
            "Network: 100 Mbps symmetric fiber. RTT to Vercel edge and to Supabase ap-south-1 are reported separately; India→database RTT must not be conflated with edge RTT (~12 ms).",
        ),
        (
            "no data loss was observed during testing",
            "autosave uses 500 ms debounce and optimistic concurrency (files.version); concurrent live editing is not supported (see Limitations)",
        ),
        (
            "formal workspace-scoped access control consistent with the RBAC model",
            "formal workspace-scoped ACL (owner/collaborator)",
        ),
        (
            "can systematically evolve from a prototype to a production-ready system",
            "can systematically evolve from a prototype to a deployable cloud-native prototype",
        ),
        (
            "First, performance measurements were conducted at a simulated concurrency of 10 users; behavior under higher loads (100+ concurrent users) remains uncharacterized. Second, the absence of WebSocket-based real-time collaboration means that concurrent editing conflicts were not evaluated. Third, the test hardware represents mid-range consumer-grade specifications; results on cloud VM instances with different memory and I/O profiles may differ.",
            "Limitations: (1) No CRDT/OT real-time co-editing—only debounced autosave with version-based OCC. (2) ACL is not enterprise RBAC. (3) PostgreSQL RLS is not enabled. (4) JWT sessions lack server-side revocation. (5) Partial Prerendering, streaming SSR, and Edge runtime were not benchmarked.",
        ),
        (
            "PostgreSQL 16 (Supabase, eu-central-1) with PgBouncer",
            "PostgreSQL 16 (Supabase ap-south-1, transaction pooler aws-1-ap-south-1)",
        ),
        (
            "Vercel (edge PoP varies by client location)",
            PROD_URL,
        ),
    ]

    total = 0
    for old, new in replacements:
        total += replace_all(doc, old, new)

    table_iv = [
        "Table IV — HTTP load test (production build, GET /, localhost, constant-arrival-rate).",
        "80 RPS: p50=5 ms, p95=8 ms, p99=23 ms, mean=6 ms, σ=5 ms, errors=0%.",
        "100 RPS: p50=4 ms, p95=7 ms, p99=19 ms, mean=5 ms, σ=3 ms, errors=0%.",
        "200 RPS: p50=3 ms, p95=7 ms, p99=25 ms, mean=4 ms, σ=5 ms, errors=0%.",
        "500 RPS: dev host saturated (24.4% errors)—not a production capacity claim.",
        "Reference Next.js @ 80 RPS (localhost): p50=2 ms, p95=3 ms, p99=4 ms.",
        "Evidence: benchmarks/results/production-suite-*.json",
    ]
    insert_after_heading(doc, "C. Functional Validation", table_iv)

    table_vi = [
        "Table VI — Production WAN load test (GET /, deployed application).",
        f"Deployment URL: {PROD_URL}",
        "Region: Vercel serverless + Supabase ap-south-1 (Mumbai) via transaction pooler.",
        "80 RPS, 30 s: p50=52 ms, p95=118 ms, p99=290 ms, mean=63 ms, σ=49 ms, errors=0%.",
        "Evidence: benchmarks/results/load-rps80-1779211899204.json",
        "Functional validation: OAuth signup/login, workspace ACL, autosave with OCC verified on production.",
    ]
    insert_after_heading(doc, "Table IV", table_vi)

    deployment = [
        "Production deployment (May 2026):",
        f"Live URL: {PROD_URL}",
        "Database: Supabase PostgreSQL (ap-south-1); serverless connections use aws-1-ap-south-1.pooler.supabase.com:6543.",
        "Auth: NextAuth.js (Google OAuth + credentials), JWT sessions, edge-safe middleware (config/auth.config.ts).",
    ]
    insert_after_heading(doc, "E. Deployment", deployment)

    appendix = [
        "Appendix A — Implementation artifacts (repository: github.com/YaGYa07/Scribeflow).",
        "Schema: lib/db/schema/app.ts · Auth: config/auth-providers.ts, lib/auth.ts · ACL: lib/db/queries/workspace.ts · DB pooler (Vercel): lib/db/database-url.ts · Autosave/OCC: lib/db/queries/file.ts",
    ]
    insert_after_heading(doc, "References", appendix)

    doc.save(OUT)
    print(f"Wrote {OUT} ({total} paragraph replacements; figures preserved from FINAL)")


if __name__ == "__main__":
    main()
