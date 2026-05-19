#!/usr/bin/env python3
"""Build ScribeFlow_REVISED.docx from docs/ScribeFlow_REVISED.md (requires python-docx)."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Install: pip install python-docx")
    raise SystemExit(1)

ROOT = Path(__file__).resolve().parents[1]
md_path = ROOT / "docs" / "ScribeFlow_REVISED.md"
out_path = ROOT / "ScribeFlow_REVISED.docx"

text = md_path.read_text(encoding="utf-8")
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

for line in text.splitlines():
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        continue
    if stripped.startswith("# "):
        p = doc.add_heading(stripped[2:], level=0)
        continue
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=1)
        continue
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=2)
        continue
    doc.add_paragraph(line)

doc.save(out_path)
print(f"Wrote {out_path}")
